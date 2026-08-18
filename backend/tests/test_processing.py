"""Tests for PDF/DOCX extraction and deterministic chunking."""

from pathlib import Path

import pymupdf
import pytest
from docx import Document as DocxDocument

from app.processing import (
    DocumentChunker,
    ExtractedDocument,
    ExtractedSegment,
    UnsupportedDocumentError,
    extract_document,
)


def test_extract_pdf_preserves_page_citations(tmp_path: Path) -> None:
    pdf_path = tmp_path / "sample.PDF"
    pdf = pymupdf.open()
    pdf.set_metadata({"title": "Retrieval Test", "author": "MIR Team"})
    first_page = pdf.new_page()
    first_page.insert_text((72, 72), "Vector space retrieval on page one.")
    second_page = pdf.new_page()
    second_page.insert_text((72, 72), "BM25 ranking evidence on page two.")
    pdf.save(pdf_path)
    pdf.close()

    extracted = extract_document(pdf_path)

    assert extracted.title == "Retrieval Test"
    assert extracted.source_format == "pdf"
    assert extracted.metadata["page_count"] == 2
    assert [segment.page_number for segment in extracted.segments] == [1, 2]
    assert extracted.text[extracted.segments[1].char_start :].startswith("BM25")
    assert extracted.segments[0].char_end < extracted.segments[1].char_start


def test_extract_docx_preserves_sections_and_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.core_properties.title = "Search Notes"
    document.add_heading("Vector Model", level=1)
    document.add_paragraph("Cosine similarity ranks documents by query relevance.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Term"
    table.cell(0, 1).text = "Weight"
    table.cell(1, 0).text = "retrieval"
    table.cell(1, 1).text = "0.75"
    document.add_heading("Probabilistic Model", level=1)
    document.add_paragraph("BM25 uses term saturation and length normalization.")
    document.save(docx_path)

    extracted = extract_document(docx_path)

    assert extracted.title == "Search Notes"
    assert extracted.source_format == "docx"
    assert extracted.metadata["table_count"] == 1
    assert [segment.section_title for segment in extracted.segments] == [
        "Vector Model",
        "Probabilistic Model",
    ]
    assert "Term | Weight" in extracted.segments[0].text
    assert "retrieval | 0.75" in extracted.segments[0].text


def test_chunker_has_overlap_and_stable_offsets() -> None:
    text = "zero one two three four five six seven eight nine ten eleven"
    document = ExtractedDocument(
        title="Numbers",
        source_format="pdf",
        text=text,
        segments=(ExtractedSegment(text=text, char_start=0, page_number=4),),
    )

    chunks = DocumentChunker(chunk_size=5, chunk_overlap=2).split(document)

    assert [chunk.text for chunk in chunks] == [
        "zero one two three four",
        "three four five six seven",
        "six seven eight nine ten",
        "nine ten eleven",
    ]
    assert [chunk.position for chunk in chunks] == [0, 1, 2, 3]
    assert all(chunk.page_start == 4 and chunk.page_end == 4 for chunk in chunks)
    assert all(
        text[chunk.char_start : chunk.char_end] == chunk.text for chunk in chunks
    )


def test_chunks_can_span_page_boundaries() -> None:
    first = " ".join(f"p1w{index}" for index in range(8))
    second = " ".join(f"p2w{index}" for index in range(8))
    full_text = f"{first}\n\n{second}"
    document = ExtractedDocument(
        title="Two pages",
        source_format="pdf",
        text=full_text,
        segments=(
            ExtractedSegment(text=first, char_start=0, page_number=1),
            ExtractedSegment(text=second, char_start=len(first) + 2, page_number=2),
        ),
    )

    chunks = DocumentChunker(chunk_size=5, chunk_overlap=2).split(document)

    assert all(
        full_text[chunk.char_start : chunk.char_end] == chunk.text
        for chunk in chunks
    )
    same_page = [chunk for chunk in chunks if chunk.page_start == chunk.page_end]
    spanned = [chunk for chunk in chunks if chunk.page_start != chunk.page_end]
    assert same_page
    assert all(chunk.page_start == chunk.page_end for chunk in same_page)
    assert spanned
    assert all(
        chunk.page_start is not None
        and chunk.page_end is not None
        and chunk.page_start < chunk.page_end
        for chunk in spanned
    )
    assert spanned[0].page_start == 1
    assert spanned[0].page_end == 2


def test_short_pages_merge_into_one_spanning_chunk() -> None:
    first = "one two three four"
    second = "five six seven eight"
    full_text = f"{first}\n\n{second}"
    document = ExtractedDocument(
        title="Two short pages",
        source_format="pdf",
        text=full_text,
        segments=(
            ExtractedSegment(text=first, char_start=0, page_number=1),
            ExtractedSegment(text=second, char_start=len(first) + 2, page_number=2),
        ),
    )

    chunks = DocumentChunker(chunk_size=10, chunk_overlap=2).split(document)

    assert len(chunks) == 1
    assert chunks[0].page_start == 1
    assert chunks[0].page_end == 2
    assert chunks[0].page_start < chunks[0].page_end
    assert full_text[chunks[0].char_start : chunks[0].char_end] == chunks[0].text


def test_unsupported_file_type_is_rejected(tmp_path: Path) -> None:
    text_path = tmp_path / "notes.txt"
    text_path.write_text("Not an allowed upload.", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentError, match="Only PDF and DOCX"):
        extract_document(text_path)


@pytest.mark.parametrize(
    ("chunk_size", "overlap"),
    [(0, 0), (10, -1), (10, 10), (10, 11)],
)
def test_invalid_chunk_configuration_is_rejected(chunk_size: int, overlap: int) -> None:
    with pytest.raises(ValueError):
        DocumentChunker(chunk_size=chunk_size, chunk_overlap=overlap)
